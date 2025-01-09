import streamlit as st
from langchain_core.prompts import PromptTemplate
from langchain_core.output_parsers import JsonOutputParser
from langchain_google_genai import ChatGoogleGenerativeAI
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from io import BytesIO
import pandas as pd
import os
import getpass
import pypandoc
import tempfile

def gemini_model():

    os.environ["GOOGLE_API_KEY"] = st.secrets['gemini']['api_key']
    model = ChatGoogleGenerativeAI(model='gemini-1.5-pro-002', temperature=0)

    return model

def data_monetization_strategy_framework(model, strategic_choices, business_strategy, business_goals_score, human_resources_score, data_technology_score, aidriven_analytics_score, data_management_score, metrics_score, monitoring_score):
    
    print('Data monetization strategy framework is processing . . .')
    
    data_monetize_prompt = """
    Persona

    You are a financial expert specializing in data monetization strategies for virtual banks in Thailand.
    You will act as an expert system, generating a comprehensive data monetization strategy framework based on user inputs.
    Your responses must be based on the definition of the "Data Monetization Strategy Framework for Virtual Banks in Thailand" provided.
    Your responses must be accurate, detailed, and comply with the purpose of the Bank of Thailand.
    Generate a complete Data Monetization Strategy Framework with recommedation based on the provided Business Strategy and Data Organization Readiness Assessment Score (Score: 0-5).
    Show the score in your recommendation text.
    The score of the Data Organization Assessment reflects the organization's readiness, which means you will provide suggestions with varying intensity based on the user's assessment level.
    Follow the Framework Structure. Ensure recommendations are realistic, actionable, and comply with Thai regulations. 
    Provide 3 use cases that related to your recommendation.
    Format the Expert System Reccomendation, Use Case as a JSON with 'expert_system_recommendation', 'use_case' as a key.
    ####

    Data Monetization Strategy Framework for Virtual Banks in Thailand

    Data monetization refers to the process of extracting economic value from data assets by leveraging them to generate new revenue streams, improve operational efficiency, and support strategic decision-making.
    Definition of Data Monetization Strategy Framework for Virtual Banks in Thailand, its contain 2 elements Strategic choices framework, Data organization readiness framework.

    - The strategic choices conceptual framework serves as a guide for organizations to balance customer experience and profit generation across four strategic quadrants: Foundation, Experience, Revenue, and Optimize. Its objective is to help businesses strategically prioritize efforts in establishing foundational digital financial products and services, strengthening customer relationships, driving revenue growth, and optimizing both profit and customer experience through data-driven personalization and advanced AI-powered services.
    1)	Foundation: Focuses on the development of essential digital infrastructure and foundational capabilities. It involves establishing robust data infrastructure, offering basic digital services, and implementing customer loyalty programs to foster engagement and enable future scalability.
    2)	Experience: Prioritizes enhancing the customer experience by building trust and fostering organic growth. It includes offering tools such as free financial planning and advisory services, focusing on strengthening brand value rather than immediate profitability.
    3)	Revenue: Focuses on implementing aggressive monetization strategies targeting existing customer bases. Strategies include driving revenue through cross-selling, upselling, and providing bundled or high-value financial products, potentially at the cost of customer experience.
    4)	Optimize: Seeks to achieve a balance between profitability and customer experience through advanced data-driven personalization and AI technologies. It emphasizes delivering highly customized financial solutions, leveraging AI for dynamic pricing and tailored customer interactions to maximize both revenue and customer satisfaction.

    - The data monetization strategy framework aims to assess the readiness of organizations by evaluating their alignment with business goals, human resources, data technology, and AI-driven analytics, while ensuring robust data management, metrics, and monitoring to maximize data value and drive strategic outcomes.
    The data monetization strategy framework consists of seven major components:
    1)	Business Goals: Define clear business objectives to align data monetization initiatives with organizational strategy. Establish a compelling value proposition to deliver measurable outcomes and drive growth.
    2)	Human Resources: Evaluate workforce capabilities, identifying skill gaps and focusing on upskilling or reskilling employees to enable seamless adoption of data-driven strategies.
    3)	Data Technology: Assess the organization’s infrastructure, including data architecture, platforms, and cloud solutions, ensuring flexibility, scalability and readiness to support advanced data monetization initiatives.
    4)	AI-Driven Analytics: Prepare the organization to harness AI and Generative AI for key operations such as acquisition campaigns, underwriting, and personalized services. Evaluate readiness to implement tailored marketing strategies and 24/7 customer support solutions like chatbots and voice bots.
    5)	Data Management: Review governance practices to ensure robust frameworks for data quality, security, compliance, accessibility, integration, and lifecycle management.
    6)	Metrics: Assess key performance indicators such as data availability and quality (e.g., accuracy), revenue growth (% uplift), customer satisfaction (CSAT), operational performance (e.g., adherence to SLAs), and cost efficiency (e.g., RODI) to evaluate success and pinpoint areas for improvement.
    7)	Monitoring: Establish systems to continuously monitor alignment with business goals e.g. dashboard, ensuring ongoing improvement and accountability in data monetization efforts.

    ####

    Bank of Thailand

    - Objective
    To introduce virtual banks in Thailand, aiming to improve financial access for underserved and unserved segments, offer better customer experiences through digital banking, and stimulate healthy competition while ensuring financial stability and sustainability.

    - Do
    Offer full-range, seamless digital banking services tailored to diverse customer needs.
    Focus on underserved segments like SMEs and low-income groups.
    Comply with regulations, maintain sound governance, and manage risks proportionately.
    Promote data portability and fair competition aligned with Open Data principles.
    Adhere to restricted operational conditions in the first 3-5 years, ensuring readiness and sustainability.

    - Do Not
    Engage in unsustainable business models or aggressive pricing that risks financial stability.
    Promote excessive debt through irresponsible lending practices.
    Give preferential treatment to related parties or abuse dominant market positions.
    Obstruct customer data portability or hinder competition.

    Example

    - Strategy
    Strategic Choices: Revenue
    Business Stragegy: Our virtual bank targets young professionals in Bangkok offering low-cost, mobile-first banking services. Competitors include traditional banks and other fintech players. Our growth objective is 5% market share within three years.

    - Data Organization Readiness Assessment Score
    Business Goals Score: 4
    Human Resources Score: 5
    Data Technology Score: 3
    AI-Driven Analytics Score: 2
    Data Management Score: 2
    Metrics Score: 1
    Monitoring Score: 1

    - Expert System Reccomendation

    "Business Goals": 
        "Recommendation": "Continue focusing on the target market of young professionals in Bangkok.  However, augment market research with advanced analytics to identify underserved sub-segments and refine your value proposition.  Explore partnerships with complementary businesses (e.g., e-commerce platforms) to increase reach and customer acquisition.",
        "Example": "Invest in a customer segmentation study using AI to identify specific needs and preferences within the target demographic. This will inform product development and marketing campaigns.",
        "KPIs": ["Customer acquisition cost (CAC) from digital channels", "Conversion rate from marketing campaigns", "Customer lifetime value (CLTV) within the target segment", "Market share growth within target demographic"]
    ,
    "Human Resources": 
        "Recommendation": "Maintain the current high level of staffing and continue investing in upskilling programs.  Focus on recruiting specialized talent (e.g., data engineers specializing in cloud infrastructure, MLOps engineers).",
        "Example": "Partner with universities to offer internships and recruit graduating students. Create mentorship programs for junior data scientists and engineers.",
        "KPIs": ["Number of data scientists and engineers hired", "Employee satisfaction scores related to training and development", "Employee retention rate", "Average time to fill open data roles"]
    ,
    "Data Technology": 
        "Recommendation": "Prioritize database infrastructure upgrades to meet the anticipated transaction volume.  Accelerate cloud migration to ensure scalability and flexibility. Invest in data integration tools to streamline data flows.",
        "Example": "Implement a new, high-performance database solution (e.g., cloud-based database service) to handle the projected transaction volume.  Migrate remaining legacy applications to the cloud.",
        "KPIs": ["Database transaction processing speed (TPS)", "System uptime", "Data integration success rate", "Cloud migration completion percentage"]
    ,
    "AI-Driven Analytics": 
        "Recommendation": "Invest in more advanced AI/ML capabilities for personalized financial products and services.  Focus on building robust AI models for risk management and fraud detection. Address data quality issues hampering effective AI model training.",
        "Example": "Develop AI-powered personalized financial recommendations based on customer transaction patterns and lifestyle data. Implement real-time fraud detection using machine learning.",
        "KPIs": ["Accuracy of fraud detection models", "Precision and recall of recommendation models", "Reduction in operational costs due to AI-driven automation", "Number of AI/ML models deployed"]
    ,
    "Data Management": 
        "Recommendation": "Enforce the data governance framework more strictly.  Invest in data quality tools and processes.  Develop a comprehensive data catalog to improve data discoverability and understanding.",
        "Example": "Implement automated data quality checks and alerts. Establish data stewardship roles within different departments. Create a centralized, searchable data catalog with detailed metadata.",
        "KPIs": ["Data quality score (percentage of data sources compliant with standards)", "Time to data discovery", "Number of data quality issues resolved", "Number of data sources documented in the catalog"]
    ,
    "Metrics": 
        "Recommendation": "Implement a comprehensive set of data-specific KPIs, covering data quality, system performance, and security.  Regularly review and refine these metrics to ensure they are aligned with business goals.",
        "Example": "Track data accuracy, completeness, and consistency. Monitor system latency and error rates. Measure the frequency and impact of data breaches.",
        "KPIs": ["Data quality score", "System uptime", "Data breach frequency and impact", "Time to insight from data analysis"]
    ,
    "Monitoring": 
        "Recommendation": "Enhance real-time monitoring capabilities and establish clear incident response protocols.  Implement automated alerts and escalation procedures.  Increase the frequency of data audits.",
        "Example": "Integrate a centralized monitoring dashboard to track key metrics. Develop a detailed incident response plan with clearly defined roles and responsibilities.  Conduct quarterly data audits.",
        "KPIs": ["Mean time to detection (MTTD) of security incidents", "Mean time to resolution (MTTR) of security incidents", "Number of security incidents", "Frequency of data audits"]

    - Use cases

        "Topic": "Targeted Marketing Campaign Leveraging Social Media and Influencers",
        "Implementation": "A targeted social media campaign on platforms popular with young professionals in Bangkok (e.g., Instagram, TikTok, Facebook) would be launched. This would involve working with relevant influencers to promote the bank's low fees, user-friendly mobile app, and other key value propositions. The campaign would focus on showcasing the convenience and affordability of the bank's services compared to traditional banks and other fintech competitors. Personalized ads based on user demographics and online behavior would be used to maximize reach and engagement.",
        "Metrics": "Number of new customer accounts opened, cost per acquisition (CPA), click-through rate (CTR) on ads, engagement rate on social media posts, brand mentions on social media."
    ,
        "Topic": "Increase customer engagement and savings deposits by 20% within one year.",
        "Implementation": "A gamified savings feature would be added to the mobile banking app. This feature could involve rewarding users for consistent saving behavior with virtual badges, points, or other incentives. It could also incorporate elements of friendly competition or challenges among users to encourage participation. The gamification would be designed to make saving money more fun and engaging for young professionals.",
        "Metrics": "Number of users actively using the gamified savings feature, average savings balance per user, user engagement with the gamified feature (e.g., time spent, frequency of interaction), growth rate of savings deposits."
    ,
        "Topic": " Expand customer base and increase transaction volume by 15% within two years.",
        "Implementation": "A strategic partnership with a leading e-commerce platform in Thailand would be established to offer seamless payment integration within the platform. Customers could use the virtual bank's mobile app to make payments directly on the e-commerce site, benefiting from the bank's low transaction fees and convenient mobile interface. This would expose the bank to a large and relevant customer base already using the e-commerce platform.",
        "Metrics": "Number of transactions processed through the e-commerce platform integration, increase in transaction volume, growth in the number of active users, customer acquisition cost through the partnership, customer satisfaction with the integrated payment system."
    
    
    ####

    - Strategy
    Strategic Choices: {strategic_choices}
    Business Stragegy: {business_strategy}

    - Data Organization Readiness Assessment Score
    Business Goals Score: {business_goals_score}
    Human Resources Score: {human_resources_score}
    Data Technology Score: {data_technology_score}
    AI-Driven Analytics Score: {aidriven_analytics_score}
    Data Management Score: {data_management_score}
    Metrics Score: {metrics_score}
    Monitoring Score: {monitoring_score}

    - Expert System Reccomendation
        Business Goals: 
            Recommendation: 
            Example:
            KPIs:
        Human Resources: 
            Recommendation: 
            Example:
            KPIs:
        Data Technology: 
            Recommendation: 
            Example:
            KPIs:
        AI-Driven Analytics:
            Recommendation: 
            Example:
            KPIs:
        Data Management:
            Recommendation: 
            Example:
            KPIs:
        Metrics:
            Recommendation: 
            Example:
            KPIs:
        Monitoring:
            Recommendation: 
            Example:
            KPIs:
            
    - Use Case

        Topic:
        Implementation:
        Metrics:
    ,
        Topic:
        Implementation:
        Metrics:
    ,
        Topic:
        Implementation:
        Metrics:

    """ 
    
    parser = JsonOutputParser()
    data_monetize_prompt_template = PromptTemplate(template=data_monetize_prompt, input_variables=['strategic_choices',
                                                                            'business_strategy',
                                                                            'business_goals_score',
                                                                            'human_resources_score',
                                                                            'data_technology_score',
                                                                            'aidriven_analytics_score',
                                                                            'data_management_score',
                                                                            'metrics_score',
                                                                            'monitoring_score'])
    data_monetize_chain = data_monetize_prompt_template | model | parser
    data_monetizen_result = data_monetize_chain.invoke({   "strategic_choices": strategic_choices,
                                                            "business_strategy": business_strategy,
                                                            "business_goals_score": business_goals_score,
                                                            "human_resources_score": human_resources_score,
                                                            "data_technology_score": data_technology_score,
                                                            "aidriven_analytics_score": aidriven_analytics_score,
                                                            "data_management_score": data_management_score,
                                                            "metrics_score": metrics_score,
                                                            "monitoring_score": monitoring_score})
    
    return data_monetizen_result

def data_architecture_framework(model, data_monetization_recommendation_information, status, business_strategy):
    
    print('Data architecture framework is processing . . .')

    data_arch_prompt = """
    Persona

    You are a financial expert specializing in data monetization strategies for virtual banks in Thailand.
    You will act as an expert system, generating a comprehensive data architecture framework based on user inputs.
    Your responses must be based on the definition of the "Data Architecture Framework for Virtual Banks in Thailand" provided.
    Your responses must be accurate, detailed, and comply with the purpose of the Bank of Thailand.
    Generate a complete Data Architecture Framework with recommedation based information provided and the status NTB or ETB founder.
    Your reponse must align with the provided information which contain the recommendation, example, and KPIs from Business Goals, Data Technology, AI-Driven Analytics, Data Management, Metrics, and Monitoring.
    Ensure your reponse for each topic of Data Architecture Reccomedation at the recommendation, example, and KPIs must significantly align with business strategy from end-to-end escepcially the use case at the data monetization opponity part.
    Follow the Framework Structure. Ensure recommendations are realistic, actionable, and comply with Thai regulations. 
    Format the Data Architecture Reccommendation as a JSON with 'data_architecture_recommendation' as a key.
    ####

    Data Architecture Framework for Virtual Banks in Thailand

    - The data architecture framework serves as a critical follow-up to the data monetization strategy framework, providing the necessary technical and operational foundation to implement monetization strategies effectively. After assessing an organization’s readiness using the data monetization strategy framework, this framework ensures that the technical infrastructure, data processes, and analytical capabilities are aligned with business goals for value creation.
    The data architecture requirements for New-to-Bank (NTB) and Existing-to-Bank (ETB) founders differ significantly, reflecting their unique challenges and advantages. NTB founders, often originating from industries outside traditional banking, must prioritize the development of scalable and flexible systems that foster innovation and adaptability to meet market demands. In contrast, ETB founders, with extensive banking expertise and access to legacy data, face the challenge of modernizing existing systems while ensuring seamless integration into virtual banking platforms. A thoughtfully designed data architecture framework is essential to addressing these distinct needs, enabling both NTB and ETB founders to establish competitive and successful virtual banks.
    1)	Data Infrastructure: Builds scalable, modern infrastructure using cloud-native solutions to ensure flexibility and future growth. Implements robust cybersecurity measures to protect customer data and maintain trust, which is critical for both acquiring new customers and serving existing ones.
    2)	Data Sources:
        a)	New to Bank: Leverages external customer data from other businesses to gain insights into customer behavior and preferences.
        b)	Existing to Bank: Combines traditional banking data (e.g., deposits, loans, and investments) to create a unified, comprehensive customer profile.
    3)	Data Collection:
        a)	New to Bank: Implement digital-first onboarding processes, loyalty programs, and social media integration to gather initial customer data and build a customer base.
        b)	Existing to Bank: Enhance existing collection mechanisms with gamified financial tools, digital-native interactions, and updated consent management to align with digital-first practices.
    4)	Data Ingestion:
        a)	New to Bank: Develop and implement modern ETL pipelines for seamless real-time and batch data ingestion from diverse digital sources.
        b)	Existing to Bank: Adapt legacy ETL systems to handle real-time data ingestion while ensuring compatibility with existing systems and processes.
    5)	Data Storage:
        a)	New to Bank: Set up cloud-based data lakes and warehouses to store data efficiently and facilitate advanced analytics.
        b)	Existing to Bank: Migrate from traditional storage systems to cloud or hybrid solutions, balancing cost, scalability, and the performance of existing operations.
    6)	AI-Driven Analytics:
        a)	New to Bank: Rapidly deploy AI/GenAI solutions for personalized services, enabling competitive differentiation from traditional banking models.
        b)	Existing to Bank: Integrate AI into their current analytics systems and develop GenAI use cases. Create innovative customer services (e.g., 24/7 chat support powered by GenAI).
    7)	Monetization Opportunities:
        a)	New to Bank: Focus on creating innovative financial products like personalized financial advice and dynamic pricing models to attract and retain customers.
        b)	Existing to Bank: Leverage existing customer relationships to cross-sell and upsell digital services, driving revenue while optimizing operations to reduce costs.
    ####

    Example
    
    Business Strategy: Our virtual bank targets young professionals in Bangkok offering low-cost, mobile-first banking services. Competitors include traditional banks and other fintech players. Our growth objective is 5% market share within three years.

    Information: 
    
    "Business Goals": 
    "Recommendation": "Continue focusing on the target market of young professionals in Bangkok.  However, augment market research with advanced analytics to identify underserved sub-segments and refine your value proposition.  Explore partnerships with complementary businesses (e.g., e-commerce platforms) to increase reach and customer acquisition.",
    "Example": "Invest in a customer segmentation study using AI to identify specific needs and preferences within the target demographic. This will inform product development and marketing campaigns.",
    "KPIs": ["Customer acquisition cost (CAC) from digital channels", "Conversion rate from marketing campaigns", "Customer lifetime value (CLTV) within the target segment", "Market share growth within target demographic"]
    ,
    "Human Resources": 
    "Recommendation": "Maintain the current high level of staffing and continue investing in upskilling programs.  Focus on recruiting specialized talent (e.g., data engineers specializing in cloud infrastructure, MLOps engineers).",
    "Example": "Partner with universities to offer internships and recruit graduating students. Create mentorship programs for junior data scientists and engineers.",
    "KPIs": ["Number of data scientists and engineers hired", "Employee satisfaction scores related to training and development", "Employee retention rate", "Average time to fill open data roles"]
    ,
    "Data Technology": 
    "Recommendation": "Prioritize database infrastructure upgrades to meet the anticipated transaction volume.  Accelerate cloud migration to ensure scalability and flexibility. Invest in data integration tools to streamline data flows.",
    "Example": "Implement a new, high-performance database solution (e.g., cloud-based database service) to handle the projected transaction volume.  Migrate remaining legacy applications to the cloud.",
    "KPIs": ["Database transaction processing speed (TPS)", "System uptime", "Data integration success rate", "Cloud migration completion percentage"]
    ,
    "AI-Driven Analytics": 
    "Recommendation": "Invest in more advanced AI/ML capabilities for personalized financial products and services.  Focus on building robust AI models for risk management and fraud detection. Address data quality issues hampering effective AI model training.",
    "Example": "Develop AI-powered personalized financial recommendations based on customer transaction patterns and lifestyle data. Implement real-time fraud detection using machine learning.",
    "KPIs": ["Accuracy of fraud detection models", "Precision and recall of recommendation models", "Reduction in operational costs due to AI-driven automation", "Number of AI/ML models deployed"]
    ,
    "Data Management": 
    "Recommendation": "Enforce the data governance framework more strictly.  Invest in data quality tools and processes.  Develop a comprehensive data catalog to improve data discoverability and understanding.",
    "Example": "Implement automated data quality checks and alerts. Establish data stewardship roles within different departments. Create a centralized, searchable data catalog with detailed metadata.",
    "KPIs": ["Data quality score (percentage of data sources compliant with standards)", "Time to data discovery", "Number of data quality issues resolved", "Number of data sources documented in the catalog"]
    ,
    "Metrics": 
    "Recommendation": "Implement a comprehensive set of data-specific KPIs, covering data quality, system performance, and security.  Regularly review and refine these metrics to ensure they are aligned with business goals.",
    "Example": "Track data accuracy, completeness, and consistency. Monitor system latency and error rates. Measure the frequency and impact of data breaches.",
    "KPIs": ["Data quality score", "System uptime", "Data breach frequency and impact", "Time to insight from data analysis"]
    ,
    "Monitoring": 
    "Recommendation": "Enhance real-time monitoring capabilities and establish clear incident response protocols.  Implement automated alerts and escalation procedures.  Increase the frequency of data audits.",
    "Example": "Integrate a centralized monitoring dashboard to track key metrics. Develop a detailed incident response plan with clearly defined roles and responsibilities.  Conduct quarterly data audits.",
    "KPIs": ["Mean time to detection (MTTD) of security incidents", "Mean time to resolution (MTTR) of security incidents", "Number of security incidents", "Frequency of data audits"]

    Status: Existing to Bank

    - Data Architecture Reccomedation

    "Data Infrastructure": 
        "Recommendation": "Migrate critical systems to a cloud-based infrastructure, ensuring high availability and scalability to handle increased transaction volumes and user interactions. Implement robust security measures to protect sensitive customer data in compliance with Thai regulations.",
        "Example": "Migrate core banking systems and databases to a cloud provider like AWS or Google Cloud. Implement multi-factor authentication, data encryption at rest and in transit, and regular security audits.",
        "KPIs": ["System uptime", "Transaction processing speed", "Mean Time To Recovery (MTTR) for system outages", "Number of security incidents"]
        ,
    "Data Sources": 
        "Recommendation": "Integrate existing customer data from various internal systems (e.g., core banking, CRM) into a unified customer data platform (CDP). Enrich this data with external sources to create a comprehensive customer profile.",
        "Example": "Implement a CDP to consolidate customer information from core banking systems, CRM, and transaction history. Enrich the data with external sources like credit bureaus and demographic data providers, respecting privacy regulations.",
        "KPIs": ["Completeness of customer profiles", "Accuracy of customer data", "Data integration success rate", "Number of data sources integrated"]
        ,
    "Data Collection": 
        "Recommendation": "Enhance existing data collection methods with digital-first approaches, such as in-app surveys, personalized notifications, and gamified financial tools.  Ensure all data collection practices are compliant with the PDPA.",
        "Example": "Implement in-app surveys to gather customer feedback on app usability and features.  Use personalized notifications to provide relevant financial advice and product recommendations.  Introduce a gamified savings feature to encourage active engagement and data collection.",
        "KPIs": ["Customer response rate to surveys", "Number of users engaging with personalized notifications", "Data collection compliance with PDPA", "Growth of customer interactions through digital channels"]
    ,
    "Data Ingestion": 
        "Recommendation": "Modernize existing ETL processes to enable real-time data ingestion and processing.  Implement robust data quality checks at the ingestion stage to ensure data accuracy and consistency.",
        "Example": "Replace outdated ETL tools with a cloud-based ETL service that supports real-time data pipelines.  Implement automated data quality checks and validation rules during the data ingestion process.",
        "KPIs": ["Data ingestion speed", "Data quality score at ingestion", "Number of data ingestion errors", "Time to integrate new data sources"]
    ,
    "Data Storage": 
        "Recommendation": "Migrate from legacy storage systems to a cloud-based data lake or warehouse to support advanced analytics and scalable data storage.  Implement data governance policies to ensure data security and compliance.",
        "Example": "Migrate data to a cloud-based data warehouse such as Snowflake or Google BigQuery.  Establish data access controls and implement data encryption.",
        "KPIs": ["Data storage cost", "Data accessibility", "Data retrieval time", "Data security incidents"]
    ,
    "AI-Driven Analytics": 
        "Recommendation": "Integrate AI/ML models into existing analytics systems to enhance customer segmentation, personalize financial advice, and improve fraud detection. Focus on GenAI for proactive customer support and improved efficiency.",
        "Example": "Develop AI-powered customer segmentation models to identify specific needs and preferences.  Build AI models to predict customer behavior and proactively offer relevant financial products.  Implement a GenAI-powered chatbot for customer support.",
        "KPIs": ["Accuracy of customer segmentation models", "Effectiveness of personalized recommendations", "Fraud detection rate", "Reduction in customer service costs"]
    ,
    "Monetization Opportunities":
        "Recommendation": "Leverage existing customer relationships and data insights to offer personalized financial products and services, enhancing revenue streams. Explore data licensing options while maintaining strict data privacy compliance.",
        "Example": "Offer personalized investment recommendations based on customer risk profiles and financial goals.  Develop targeted marketing campaigns based on customer segments.  Explore opportunities to license anonymized and aggregated customer data while strictly complying with PDPA.",
        "KPIs": ["Revenue generated from personalized products", "Number of customers using personalized services", "Revenue from data licensing", "Customer churn rate"]
    ###

    Business Strategy: {business_strategy}

    Information: {data_monetization_recommendation_information}

    Status: {status}

    - Data Architecture Reccomedation

    Data Infrastructure:
        Recommendation: 
        Example:
        KPIs:
    Data Sources:
        Recommendation: 
        Example:
        KPIs:
    Data Collection:
        Recommendation: 
        Example:
        KPIs:
    Data Ingestion:
        Recommendation: 
        Example:
        KPIs:
    Data Storage:
        Recommendation: 
        Example:
        KPIs:
    AI-Driven Analytics:
        Recommendation: 
        Example:
        KPIs:
    Monetization Opportunities:
        Recommendation: 
        Example:
        KPIs:

    """
    
    parser = JsonOutputParser()
    data_arch_prompt_template = PromptTemplate(template=data_arch_prompt, input_variables=['data_monetization_recommendation_information', 'status', 'business_strategy'])
    data_arch_chain = data_arch_prompt_template | model | parser
    data_arch_result = data_arch_chain.invoke({"data_monetization_recommendation_information": data_monetization_recommendation_information, 
                                                "status": status,
                                                "business_strategy": business_strategy})
    
    return data_arch_result

def save_docs(data_monetize_result_df, data_arch_result_df, use_case_df):
    
    doc = Document()
    
    title1 = doc.add_heading('Data Monetization Strategy Framework for Virtual Banks in Thailand', level=1)
    title1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    run = title1.runs[0]
    run.font.size = Pt(18) 
    run.font.name = 'Arial'
    run.font.color.rgb = RGBColor(2,2,2)
    
    title2 = doc.add_heading('Data Monetization Strategy Framework', level=2)
    title2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    run = title2.runs[0]
    run.font.size = Pt(12) 
    run.font.name = 'Arial'
    run.font.color.rgb = RGBColor(128,0,128)
    paragraph2 = doc.add_paragraph("This section provides details on the data monetization strategy framework.")
    paragraph2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph("")

    table = doc.add_table(rows=1, cols=len(data_monetize_result_df.columns) + 1)
    table.style = 'Light List'

    header_cells = table.rows[0].cells
    header_cells[0].text = 'Category'
    for i, col_name in enumerate(data_monetize_result_df.columns, start=1):
        header_cells[i].text = col_name
        
    for idx, row in data_monetize_result_df.iterrows():
        cells = table.add_row().cells
        cells[0].text = idx
        for i, value in enumerate(row):
            cells[i + 1].text = str(value)

        for cell in cells:
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            run = paragraph.runs[0]
            run.font.size = Pt(8)
            run.font.name = 'Arial'
        
    doc.add_page_break()
    
    title3 = doc.add_heading('Data Architecture Framework', level=2)
    title3.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    run = title3.runs[0]
    run.font.size = Pt(14) 
    run.font.name = 'Arial'
    run.font.color.rgb = RGBColor(128,0,128)
    paragraph3 = doc.add_paragraph("This section provides details on the data architecture framework.")
    paragraph3.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph("")

    table = doc.add_table(rows=1, cols=len(data_arch_result_df.columns) + 1)
    table.style = 'Light List'

    header_cells = table.rows[0].cells
    header_cells[0].text = 'Category'
    for i, col_name in enumerate(data_arch_result_df.columns, start=1):
        header_cells[i].text = col_name
        
    for idx, row in data_arch_result_df.iterrows():
        cells = table.add_row().cells
        cells[0].text = idx
        for i, value in enumerate(row):
            cells[i + 1].text = str(value)

        for cell in cells:
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            run = paragraph.runs[0]
            run.font.size = Pt(8)
            run.font.name = 'Arial'

    doc.add_page_break()
    
    title4 = doc.add_heading('Use Case Example', level=2)
    title4.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    run = title4.runs[0]
    run.font.size = Pt(14) 
    run.font.name = 'Arial'
    run.font.color.rgb = RGBColor(128,0,128)
    # doc.add_paragraph("")
    paragraph3 = doc.add_paragraph("This section provides examples use case.")
    paragraph3.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph("")

    table = doc.add_table(rows=1, cols=len(use_case_df.columns) + 1)
    table.style = 'Light List'

    header_cells = table.rows[0].cells
    header_cells[0].text = 'Category'
    for i, col_name in enumerate(use_case_df.columns, start=1):
        header_cells[i].text = col_name
        
    for idx, row in use_case_df.iterrows():
        cells = table.add_row().cells
        cells[0].text = idx
        for i, value in enumerate(row):
            cells[i + 1].text = str(value)

        for cell in cells:
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            run = paragraph.runs[0]
            run.font.size = Pt(8)
            run.font.name = 'Arial'

    doc_file = BytesIO()
    doc.save(doc_file)
    doc_file.seek(0)

    return doc_file

def main():

    if 'doc_content' not in st.session_state:
        st.session_state['doc_content'] = None

    st.title("Data Monetization Strategy Framework for Virtual Banks in Thailand")
    st.header(":red[Strategy]")
    strategic_choices = st.selectbox(
        "Strategic Choices",
        ["Foundation", "Experience", "Revenue", "Optimize"]
    )
    status = st.selectbox(
        "Status",
        ["New to Bank", "Existing to Bank"]
    )
    business_strategy = st.text_area("Business Strategy")

    st.header(":red[Data Organization Readiness Assessment Score (0-5)]")
    business_goals_score = st.slider("Business Goals Score", 0, 5, 2)
    human_resources_score = st.slider("Human Resources Score", 0, 5, 2)
    data_technology_score = st.slider("Data Technology Score", 0, 5, 2)
    aidriven_analytics_score = st.slider("AI-Driven Analytics Score", 0, 5, 2)
    data_management_score = st.slider("Data Management Score", 0, 5, 2)
    metrics_score = st.slider("Metrics Score", 0, 5, 2)
    monitoring_score = st.slider("Monitoring Score", 0, 5, 2)

    if st.button("Submit"):
        with st.spinner("Processing your document..."):
            model = gemini_model()
            data_monetization_strategy_framework_result = data_monetization_strategy_framework(model,
                                                                                            strategic_choices,
                                                                                            business_strategy,
                                                                                            business_goals_score,
                                                                                            human_resources_score,
                                                                                            data_technology_score,
                                                                                            aidriven_analytics_score,
                                                                                            data_management_score,
                                                                                            metrics_score,
                                                                                            monitoring_score)

            data_monetization_recommendation_information = data_monetization_strategy_framework_result['expert_system_recommendation']
            data_architecture_framework_result = data_architecture_framework(model, 
                                                                            data_monetization_recommendation_information, 
                                                                            status, 
                                                                            business_strategy)
            data_architecture_framework_information = data_architecture_framework_result['data_architecture_recommendation']
            use_case_example_information = {
                use_case['Topic']: {
                    "Use_Case": use_case["Topic"],
                    "Implementation": use_case["Implementation"],
                    "Metrics": use_case["Metrics"]
                }
                for use_case in data_monetization_strategy_framework_result['use_case']
            }

            data_monetization_strategy_framework_df = pd.DataFrame.from_dict(data_monetization_recommendation_information, orient='index')
            data_architecture_framework_df = pd.DataFrame.from_dict(data_architecture_framework_information, orient='index')
            use_case_example_df = pd.DataFrame.from_dict(use_case_example_information, orient='index')

            st.session_state['doc_content'] = save_docs(
                data_monetize_result_df=data_monetization_strategy_framework_df,
                data_arch_result_df=data_architecture_framework_df,
                use_case_df=use_case_example_df
            )

            st.success("Document processed successfully!")

        if st.session_state['doc_content']:
            st.markdown("---")
            st.markdown(
                """
                <div style="background-color: #f9f9f9; padding: 20px; border-radius: 10px; text-align: center;">
                    <h3 style="color: #4CAF50;">Your document is ready to download! 📄</h3>
                    <p style="color: #555;">Click the button below to download your customized data monetization strategy document.</p>
                </div>
                """,
                unsafe_allow_html=True
            )
            
            st.download_button(
                label="📥 Download Document",
                data=st.session_state['doc_content'],
                file_name="data_monetization_strategy_framework_docs.docx",
                key="download-button"
            )

if __name__ == '__main__':
    main()